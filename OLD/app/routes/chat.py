# /workspace/app/routes/query.py
"""
/query 用ルーター
"""

from fastapi import APIRouter, Form
from fastapi.responses import JSONResponse, StreamingResponse
from OLD.app.services.chat.handler import handle_query
import json
import os
import datetime
from typing import List, Dict
from io import BytesIO

router = APIRouter()

# 履歴ファイルのパス
HISTORY_FILE = "logs/search_history.json"

# 処理停止用のグローバル変数
_current_search_task = None
_search_cancelled = False

async def save_search_history(query: str, model_key: str, mode: str, result: Dict, processing_time: float):
    """検索履歴を保存"""
    try:
        # 履歴ディレクトリを作成
        os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
        
        # 既存の履歴を読み込み
        history = []
        if os.path.exists(HISTORY_FILE):
            try:
                with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if content:
                        history = json.loads(content)
                    else:
                        history = []
            except (json.JSONDecodeError, UnicodeDecodeError) as e:
                print(f"履歴ファイル読み込みエラー（破損ファイルを初期化）: {e}")
                history = []
        
        # 新しい履歴エントリを作成
        history_entry = {
            "id": len(history) + 1,
            "timestamp": datetime.datetime.now().isoformat(),
            "query": query,
            "model_key": model_key,
            "mode": mode,
            "processing_time": processing_time,
            "result": result
        }
        
        # 履歴に追加（最新を先頭に）
        history.insert(0, history_entry)
        
        # 履歴を最大100件に制限
        if len(history) > 100:
            history = history[:100]
        
        # ファイルに保存（原子的書き込みで破損を防止）
        import tempfile
        temp_file = HISTORY_FILE + '.tmp'
        try:
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
            # 原子的にファイルを置き換え
            os.replace(temp_file, HISTORY_FILE)
        except Exception as e:
            # 一時ファイルが残っている場合は削除
            if os.path.exists(temp_file):
                os.remove(temp_file)
            raise e
            
    except Exception as e:
        print(f"履歴保存エラー: {e}")

async def load_search_history() -> List[Dict]:
    """検索履歴を読み込み"""
    try:
        if not os.path.exists(HISTORY_FILE):
            return []
        
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"履歴読み込みエラー: {e}")
        return []

async def remove_search_history(history_id: int):
    """指定された履歴を削除"""
    try:
        history = await load_search_history()
        history = [h for h in history if h.get("id") != history_id]
        
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"履歴削除エラー: {e}")

async def clear_all_search_history():
    """全履歴を削除"""
    try:
        if os.path.exists(HISTORY_FILE):
            os.remove(HISTORY_FILE)
    except Exception as e:
        print(f"履歴全削除エラー: {e}")

@router.post("/query")
async def query_api(
    query: str        = Form(...),
    model_key: str    = Form(...),
    mode: str         = Form("チャンク統合"),
    search_limit: int = Form(10),
    min_score: float  = Form(0.0)
):
    global _current_search_task, _search_cancelled
    
    try:
        import time
        import asyncio
        
        # キャンセルフラグをリセット
        _search_cancelled = False
        
        start_time = time.time()
        
        # キャンセル可能なタスクとして実行
        _current_search_task = asyncio.create_task(
            handle_query(query, model_key, mode, search_limit, min_score)
        )
        
        result = await _current_search_task
        processing_time = round(time.time() - start_time, 2)
        
        # 履歴に保存
        await save_search_history(query, model_key, mode, result, processing_time)
        
        return result
        
    except asyncio.CancelledError:
        return JSONResponse({"error": "検索処理がキャンセルされました"}, status_code=499)
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)
    finally:
        _current_search_task = None
        _search_cancelled = False

@router.get("/history")
async def get_search_history():
    """検索履歴を取得"""
    try:
        return await load_search_history()
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.delete("/history/{history_id}")
async def delete_search_history(history_id: int):
    """指定された履歴を削除"""
    try:
        await remove_search_history(history_id)
        return {"status": "deleted"}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.delete("/history")
async def clear_search_history():
    """全履歴を削除"""
    try:
        await clear_all_search_history()
        return {"status": "cleared"}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.post("/stop_search")
async def stop_current_search():
    """現在実行中の検索処理を停止"""
    global _search_cancelled
    try:
        _search_cancelled = True
        
        # 現在のタスクがあればキャンセル
        if _current_search_task and not _current_search_task.done():
            _current_search_task.cancel()
            
        return {"status": "stopped", "message": "検索処理の停止を要求しました"}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

@router.get("/history/download/{format_type}")
async def download_search_history(format_type: str):
    """検索履歴をダウンロード"""
    try:
        history = await load_search_history()
        if not history:
            return JSONResponse({"error": "履歴がありません"}, status_code=404)
        
        # ファイル生成
        buffer = await generate_history_file(history, format_type)
        
        # MIMEタイプとファイル名を設定
        mime_map = {
            "txt": "text/plain",
            "rtf": "application/rtf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "json": "application/json"
        }
        
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        filename = f"search_history_{today_str}.{format_type}"
        
        return StreamingResponse(
            BytesIO(buffer),
            media_type=mime_map.get(format_type, "application/octet-stream"),
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)

async def generate_history_file(history: List[Dict], format_type: str) -> bytes:
    """履歴ファイルを生成"""
    if format_type == "txt":
        content = ""
        for idx, h in enumerate(history, 1):
            timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
            content += (
                f"{idx}. 質問: {h['query']}\n"
                f"モデル: {h['model_key']} | モード: {h['mode']}\n"
                f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}\n"
            )
            
            # 結果の表示
            result = h.get('result', {})
            if result.get('mode') == 'チャンク統合':
                content += f"統合回答:\n{result.get('answer', '')}\n"
                if result.get('sources'):
                    content += "参照ファイル:\n"
                    for src in result['sources']:
                        content += f"  - {src.get('file_name', '')}\n"
            else:
                content += "ファイル別結果:\n"
                for res in result.get('results', []):
                    content += f"  - {res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\n"
                    content += f"    要約: {res.get('summary', '')}\n"
            content += "\n" + "="*50 + "\n\n"
        
        return content.encode("utf-8")
    
    elif format_type == "json":
        return json.dumps(history, ensure_ascii=False, indent=2).encode("utf-8")
    
    elif format_type == "rtf":
        content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24\n"
        for idx, h in enumerate(history, 1):
            timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
            content += (
                f"\\b {idx}. 質問: {h['query']}\\b0\\line\n"
                f"モデル: {h['model_key']} | モード: {h['mode']}\\line\n"
                f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}\\line\\line\n"
            )
            
            result = h.get('result', {})
            if result.get('mode') == 'チャンク統合':
                content += f"\\i 統合回答:\\i0\\line\n{result.get('answer', '')}\\line\\line\n"
            else:
                content += "\\i ファイル別結果:\\i0\\line\n"
                for res in result.get('results', []):
                    content += f"{res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\\line\n"
            content += "\\line\n"
        content += "}"
        return content.encode("utf-8")
    
    elif format_type == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            doc.add_heading('検索履歴', 0)
            
            for idx, h in enumerate(history, 1):
                timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
                
                # 質問
                doc.add_heading(f"{idx}. 質問: {h['query']}", level=2)
                
                # メタ情報
                meta_p = doc.add_paragraph()
                meta_p.add_run(f"モデル: {h['model_key']} | モード: {h['mode']}\n")
                meta_p.add_run(f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}")
                
                # 結果
                result = h.get('result', {})
                if result.get('mode') == 'チャンク統合':
                    doc.add_heading('統合回答', level=3)
                    answer_p = doc.add_paragraph(result.get('answer', ''))
                    answer_p.style.font.size = Pt(11)
                else:
                    doc.add_heading('ファイル別結果', level=3)
                    for res in result.get('results', []):
                        file_p = doc.add_paragraph()
                        file_p.add_run(f"{res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\n").bold = True
                        file_p.add_run(f"要約: {res.get('summary', '')}")
                
                doc.add_page_break()
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            # python-docxがインストールされていない場合はテキスト形式で返す
            return await generate_history_file(history, "txt")
    
    else:
        raise ValueError(f"サポートされていない形式: {format_type}")
