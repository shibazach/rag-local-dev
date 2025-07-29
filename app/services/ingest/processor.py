# app/services/ingest/processor.py
# Ingest処理の新しい実装（共通OCRサービス使用）

import functools
import time
import unicodedata
import os
import logging
from typing import Dict, Generator, List, AsyncGenerator

from src.config import OLLAMA_MODEL
from fileio.file_embedder import embed_and_insert
from llm.chunker import split_into_chunks
from llm.refiner import (
    normalize_empty_lines,
    refine_text_with_llm,
    build_prompt,
)
from ocr.spellcheck import correct_text
from db.handler import (
    insert_file_full,
    update_file_text,
    update_file_status,
    get_file_text,
)

# 新しい共通OCRサービスをインポート
from app.services.ocr import OCREngineFactory
# マルチモーダル処理をインポート
from app.services.multimodal.processor import multimodal_processor
# 非同期キャンセル機能をインポート
from app.services.async_cancellation import cancellation_manager, OllamaCancellationHelper

# 定数
TOK_LIMIT = 8192
CHUNK_SIZE = 500
OVERLAP_SIZE = 50
LOGGER = logging.getLogger("ingest_processor")

# 英語テンプレ誤反映などに該当する典型パターン（lower比較前提）
TEMPLATE_PATTERNS = [
    "certainly", "please provide", "reformat", "i will help",
    "note that this is a translation", "based on your ocr output"
]

class IngestProcessor:
    """Ingest処理のメインクラス"""
    
    def __init__(self):
        self.ocr_factory = OCREngineFactory()
    
    async def process_file(
        self,
        *,
        file_path: str,
        file_name: str,
        index: int,
        total_files: int,
        refine_prompt_key: str,
        refine_prompt_text: str = None,
        embed_models: List[str],
        overwrite_existing: bool,
        quality_threshold: float,
        llm_timeout_sec: int,
        abort_flag: Dict[str, bool],
        ocr_engine_id: str = None,  # 新規追加: OCRエンジン指定
        ocr_settings: Dict = None,  # 新規追加: OCR設定
    ) -> AsyncGenerator[Dict, None]:
        """1ファイルを処理しイベントを yield"""
        t0 = time.perf_counter()

        # ── ① DB 仮登録 ───────────────────────────────────
        try:
            file_id = insert_file_full(file_path, "", "", 0.0)
        except Exception as exc:
            LOGGER.exception("insert_file_full")
            yield {"file": file_name, "step": "DB登録失敗", "detail": str(exc)}
            return

        yield {
            "file": file_name,
            "file_id": file_id,
            "step": "ファイル登録完了",
            "index": index,
            "total": total_files,
        }

        # ── ② テキスト抽出（新しいOCRサービス使用） ─────────────────────────────────
        yield {"file": file_name, "step": "テキスト抽出開始"}
        
        try:
            # OCR処理をジェネレータとして実行し、進捗を直接yield
            pages = []
            for event_or_result in self._extract_text_with_ocr_generator(
                file_path, 
                file_name, 
                ocr_engine_id, 
                ocr_settings
            ):
                if isinstance(event_or_result, dict) and "step" in event_or_result:
                    # 進捗イベントを即座にyield
                    yield event_or_result
                elif isinstance(event_or_result, list):
                    # 最終結果（ページリスト）を受け取り
                    pages = event_or_result
        except Exception as exc:
            LOGGER.exception("extract_text_with_ocr")
            update_file_status(file_id, status="error", note=str(exc))
            yield {"file": file_name, "step": "テキスト抽出失敗", "detail": str(exc)}
            return

        if not pages:
            update_file_status(file_id, status="error", note="no text")
            yield {"file": file_name, "step": "テキスト無し"}
            return

        yield {"file": file_name, "step": "テキスト抽出完了", "pages": len(pages)}

        # ── ③ 全文生成 → raw_text 保存 ────────────────────
        page_blocks = [
            normalize_empty_lines(correct_text(unicodedata.normalize("NFKC", p)))
            for p in pages
        ]
        full_text = "\n\n".join(page_blocks)
        
        # マルチモーダル処理を適用（PDFファイルの場合）
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            # GPU環境でのみマルチモーダル処理を実行
            from src.config import MULTIMODAL_ENABLED
            if MULTIMODAL_ENABLED:
                full_text = await self._apply_multimodal_processing(
                    full_text, file_path, file_name, 
                    progress_yield=None  # 非同期ジェネレータ内では直接yieldできないためNone
                )
            else:
                yield {"file": file_name, "step": "マルチモーダル処理をスキップ（CPU環境）"}
        
        update_file_text(file_id, raw_text=full_text)

        token_est = int(len(full_text) * 1.6)
        use_page = token_est > TOK_LIMIT
        unit_iter = enumerate(page_blocks, 1) if use_page else [("all", full_text)]

        refined_pages: List[str] = []
        file_min_score: float = 1.0

        # ── ④ LLM 整形 ─────────────────────────────────────
        # EMLファイルの場合は段落別整形処理を適用
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".eml":
            yield {"file": file_name, "step": "📧 EMLファイル段落整形モード適用"}
            
            # 引用記号の除去と段落分割
            corrected = full_text.replace(">>", "").replace("> >", "").replace("> ", "")
            paragraphs = [p.strip() for p in corrected.split("\n\n") if len(p.strip()) > 30]
            
            yield {"file": file_name, "step": f"段落分割完了: {len(paragraphs)}個の段落"}
            
            refined_parts = []
            for i, para in enumerate(paragraphs):
                if abort_flag["flag"]:
                    return
                    
                yield {"file": file_name, "step": f"📑 段落 {i+1}/{len(paragraphs)} を整形中..."}
                
                try:
                    # 段落ごとにLLM整形
                    refined, _, score, _ = refine_text_with_llm(
                        para, 
                        model=OLLAMA_MODEL,
                        force_lang=refine_prompt_key,
                        abort_flag=abort_flag
                    )
                    
                    # 不正な整形結果をチェック
                    if self._is_invalid_llm_output(refined):
                        yield {"file": file_name, "step": f"⚠️ 段落 {i+1}: 不正な整形結果をスキップ"}
                        continue
                        
                    refined_parts.append(refined)
                    file_min_score = min(file_min_score, score)
                    
                    yield {"file": file_name, "step": f"✅ 段落 {i+1}/{len(paragraphs)} 整形完了"}
                    
                except Exception as e:
                    LOGGER.warning(f"段落 {i+1} 整形失敗: {e}")
                    yield {"file": file_name, "step": f"⚠️ 段落 {i+1} 整形失敗: {e}"}
                    continue
            
            if refined_parts:
                refined_pages = ["\n\n".join(refined_parts)]
                yield {"file": file_name, "step": f"EML段落整形完了: {len(refined_parts)}段落を統合"}
            else:
                yield {"file": file_name, "step": "⚠️ 有効な段落が見つかりませんでした"}
                refined_pages = [full_text]  # フォールバック
        else:
            # 通常のファイル処理
            for pg, block in unit_iter:
                if abort_flag["flag"]:
                    return

                label = f"page{pg}" if use_page else "all"
                # 編集ペインのテキストがあればそれを使う（{TEXT}置換のみ）
                if refine_prompt_text:
                    cleaned = normalize_empty_lines(correct_text(block))
                    prompt_text = refine_prompt_text.replace("{TEXT}", cleaned)
                else:
                    prompt_text = build_prompt(block, refine_prompt_key)

                # プロンプト見出し（全文用）
                yield {"file": file_name, "step": f"使用プロンプト全文 part:{label}"}
                # プロンプト全文
                yield {
                    "file": file_name,
                    "step": "prompt_text",
                    "part": label,
                    "content": prompt_text,
                }

                # 進行中メッセージ（※廃止しない）
                yield {"file": file_name, "step": f"LLM整形中 part:{label}"}

                # 非同期キャンセル機能を使用したLLM処理
                task_id = f"llm_refine_{file_name}_{label}"
                helper = OllamaCancellationHelper(task_id)
                
                try:
                    # キャンセル可能なLLM処理を実行
                    refined, _, score, _ = await helper.run_ollama_with_cancellation(
                        lambda: refine_text_with_llm(
                            block,
                            model=OLLAMA_MODEL,
                            force_lang=refine_prompt_key,
                            abort_flag=abort_flag,
                        ),
                        timeout=llm_timeout_sec if llm_timeout_sec > 0 else 300
                    )
                    
                except asyncio.CancelledError:
                    update_file_status(file_id, status="error", note="llm cancelled")
                    yield {"file": file_name, "step": "LLM処理がキャンセルされました", "part": label}
                    return
                except TimeoutError:
                    update_file_status(file_id, status="error", note="llm timeout")
                    yield {"file": file_name, "step": "LLMタイムアウト", "part": label}
                    return
                except Exception as exc:
                    LOGGER.exception("LLM refine")
                    update_file_status(file_id, status="error", note=str(exc))
                    yield {"file": file_name, "step": "LLM整形失敗", "detail": str(exc)}
                    return

                refined_pages.append(refined)
                file_min_score = min(file_min_score, score)

                # 整形結果見出し（全文）
                yield {"file": file_name, "step": f"LLM整形結果全文 part:{label}"}
                # 整形結果全文
                yield {
                    "file": file_name,
                    "step": "refined_text",
                    "part": label,
                    "content": refined,
                }

        if abort_flag["flag"] or not refined_pages:
            return

        # ── ⑤ チャンク分割 ─────────────────────────────────
        final_text = "\n\n".join(refined_pages)
        chunks = split_into_chunks(final_text, chunk_size=CHUNK_SIZE, overlap=OVERLAP_SIZE)
        yield {"file": file_name, "step": f"チャンク分割完了 chunks={len(chunks)}"}

        # ── ⑥ ベクトル化 ───────────────────────────────────
        embed_job = functools.partial(
            embed_and_insert,
            texts=chunks,
            file_name=file_path,
            model_keys=embed_models,
            quality_score=file_min_score,
            file_id=file_id,
        )
        try:
            embed_job()
            yield {"file": file_name, "step": "ベクトル化完了"}
        except Exception as exc:
            LOGGER.exception("embed_and_insert")
            update_file_status(file_id, status="error", note=str(exc))
            yield {"file": file_name, "step": "ベクトル化失敗", "detail": str(exc)}
            return

        # ── ⑦ refined_text 保存 ────────────────────────────
        cur = get_file_text(file_id) or {}
        need_upd = (
            overwrite_existing or not cur.get("refined_text") or
            (cur.get("quality_score") or 0.0) < quality_threshold or
            file_min_score < (cur.get("quality_score") or 1.0)
        )
        if need_upd:
            update_file_text(file_id, refined_text=final_text, quality_score=file_min_score)
            yield {"file": file_name, "step": "全文保存完了（上書き実施）"}
        else:
            yield {"file": file_name, "step": "全文保存スキップ"}

        update_file_status(file_id, status="done")
        yield {
            "file": file_name,
            "step": "file_done",
            "elapsed": round(time.perf_counter() - t0, 2),
        }

    def _extract_text_with_ocr(
        self, 
        file_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None,
        progress_yield=None
    ) -> List[str]:
        """新しいOCRサービスを使用してテキスト抽出"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            return self._extract_pdf_with_ocr(file_path, file_name, ocr_engine_id, ocr_settings, progress_yield)
        elif ext == ".docx":
            return self._extract_docx_with_ocr(file_path, file_name, ocr_engine_id, ocr_settings, progress_yield)
        elif ext == ".txt":
            if progress_yield:
                progress_yield({"file": file_name, "step": "テキストファイルを処理中..."})
            return self._extract_text_from_txt(file_path)
        elif ext == ".csv":
            if progress_yield:
                progress_yield({"file": file_name, "step": "CSVファイルを処理中..."})
            return self._extract_text_from_csv(file_path)
        elif ext == ".json":
            if progress_yield:
                progress_yield({"file": file_name, "step": "JSONファイルを処理中..."})
            return self._extract_text_from_json(file_path)
        elif ext == ".eml":
            if progress_yield:
                progress_yield({"file": file_name, "step": "EMLファイルを処理中..."})
            return self._extract_text_from_eml(file_path)
        else:
            raise ValueError(f"未対応のファイル形式: {ext}")

    def _extract_pdf_with_ocr(
        self, 
        pdf_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None,
        progress_yield=None
    ) -> List[str]:
        """PDFからテキスト抽出（新しいOCRサービス使用）"""
        import fitz
        
        doc = fitz.open(pdf_path)
        text_list = []
        total_pages = len(doc)
        
        # OCRエンジンの決定
        if not ocr_engine_id:
            ocr_engine_id = self.ocr_factory.settings_manager.get_default_engine()
        
        if progress_yield:
            progress_yield({"file": file_name, "step": f"OCRエンジン: {ocr_engine_id} を使用"})

        for page_number in range(total_pages):
            if progress_yield:
                progress_yield({"file": file_name, "step": f"ページ {page_number + 1}/{total_pages} 処理中..."})

            page = doc.load_page(page_number)

            # PyMuPDF抽出（テキスト層）
            pdf_text = page.get_text().strip()

            # 新しいOCRサービスでOCR抽出
            ocr_result = self.ocr_factory.process_with_settings(
                engine_id=ocr_engine_id,
                pdf_path=pdf_path,
                page_num=page_number,
                custom_params=ocr_settings or {}
            )
            
            if ocr_result["success"]:
                ocr_text = ocr_result["text"].strip()
                if progress_yield:
                    progress_yield({
                        "file": file_name, 
                        "step": f"ページ {page_number + 1} OCR完了 (処理時間: {int(ocr_result['processing_time'])}秒)"
                    })
            else:
                ocr_text = f"OCRエラー: {ocr_result['error']}"
                if progress_yield:
                    progress_yield({
                        "file": file_name, 
                        "step": f"ページ {page_number + 1} OCRエラー: {ocr_result['error']}"
                    })

            # 構造化データの抽出
            structural_data = self._extract_structural_data(page, page_number)
            
            # タグ形式で結合して1ページ分として格納
            merged = f"<pdf_text>\n{pdf_text}\n</pdf_text>\n\n<ocr_text>\n{ocr_text}\n</ocr_text>\n\n<structural_data>\n{structural_data}\n</structural_data>"
            text_list.append(merged)

        doc.close()
        return text_list

    def _extract_structural_data(self, page, page_number: int) -> str:
        """ページから構造化データを抽出"""
        try:
            # テキストブロックの抽出
            blocks = page.get_text("dict")["blocks"]
            structural_info = []
            
            for block in blocks:
                if block['type'] == 0:  # テキストブロック
                    bbox = block['bbox']
                    block_text = ""
                    for line in block["lines"]:
                        line_text = " ".join(span["text"] for span in line["spans"])
                        block_text += line_text + "\n"
                    
                    if block_text.strip():
                        structural_info.append({
                            "type": "text_block",
                            "text": block_text.strip(),
                            "bbox": bbox,
                            "page": page_number + 1
                        })
                elif block['type'] == 1:  # 画像ブロック
                    structural_info.append({
                        "type": "image",
                        "bbox": block['bbox'],
                        "page": page_number + 1
                    })
            
            # 構造化データをJSON形式で返す
            import json
            return json.dumps(structural_info, ensure_ascii=False, indent=2)
            
        except Exception as e:
            LOGGER.warning(f"構造化データ抽出エラー: {e}")
            return "{}"

    def _extract_docx_with_ocr(
        self, 
        docx_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None,
        progress_yield=None
    ) -> List[str]:
        """Wordファイルから構造とOCRを統合して抽出"""
        import docx
        import subprocess
        import tempfile
        
        structured_text = []

        if progress_yield:
            progress_yield({"file": file_name, "step": "Word文書の構造を抽出中..."})

        # 表や段落など、論理構造ベースで抽出
        doc = docx.Document(docx_path)
        structured_text.extend([p.text for p in doc.paragraphs if p.text])

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if row_text:
                    structured_text.append(" | ".join(row_text))

        # OCRベース（LibreOffice経由でPDF化 → 新しいOCRサービス使用）
        with tempfile.TemporaryDirectory() as tmpdir:
            if progress_yield:
                progress_yield({"file": file_name, "step": "Word文書をPDFに変換中..."})
                
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", tmpdir, docx_path
                ], check=True)
            except subprocess.CalledProcessError:
                raise RuntimeError("LibreOfficeによるPDF変換に失敗しました")

            basename = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(tmpdir, f"{basename}.pdf")
            if os.path.exists(pdf_path):
                ocr_text = self._extract_pdf_with_ocr(pdf_path, file_name, ocr_engine_id, ocr_settings, progress_yield)
            else:
                raise FileNotFoundError(f"PDF変換後ファイルが見つかりません: {pdf_path}")

        # 統合：タグを付けて LLM 整形に渡す
        structured_content = "\n".join(structured_text)
        ocr_content = "\n".join(ocr_text)
        
        # 両内容を統合して1つのテキストとして返す
        combined_text = f"<structured_text>\n{structured_content}\n</structured_text>\n\n<ocr_text>\n{ocr_content}\n</ocr_text>"
        return [combined_text]

    def _extract_text_from_txt(self, txt_path: str) -> List[str]:
        """TXTファイルからテキスト抽出"""
        with open(txt_path, encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]

    def _extract_text_from_csv(self, csv_path: str) -> List[str]:
        """CSVファイルからテキスト抽出"""
        import csv
        results = []
        try:
            with open(csv_path, encoding="utf-8") as f:
                reader = csv.DictReader(f)
                if "text" not in reader.fieldnames:
                    raise ValueError("CSVに'text'カラムが必要です")
                for row in reader:
                    if row.get("text"):
                        results.append(row["text"].strip())
            return results
        except Exception as e:
            os.makedirs("logs", exist_ok=True)
            log_path = os.path.join("logs", "invalid_csv_log.txt")
            from datetime import datetime
            now = datetime.utcnow().isoformat() + "Z"
            with open(log_path, "a", encoding="utf-8") as logf:
                logf.write(f"{now}: Skipped {csv_path} - {e}\n")
            return []

    def _extract_text_from_json(self, json_path: str) -> List[str]:
        """JSONファイルからテキスト抽出"""
        import json
        with open(json_path, encoding="utf-8") as f:
            data = json.load(f)
            if not isinstance(data, list):
                raise ValueError("JSONはリスト形式である必要があります")
            return [item["text"] for item in data if "text" in item]

    def _extract_text_from_eml(self, file_path: str) -> List[str]:
        """EMLファイルからテキスト抽出"""
        from email import policy
        from email.parser import BytesParser

        with open(file_path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        text_blocks = []
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                content = part.get_content()
                lines = content.splitlines()
                # 引用行（>>で始まる行）は除去
                lines = [line for line in lines if not line.strip().startswith(">>")]
                filtered = "\n".join(lines).strip()
                if filtered:
                    text_blocks.append(filtered)
        return text_blocks

    def _is_invalid_llm_output(self, text: str) -> bool:
        """LLM整形後の出力がテンプレ・英語・無意味などの不正内容かを判定"""
        try:
            from langdetect import detect
        except ImportError:
            # langdetectがない場合は基本チェックのみ
            if len(text.strip()) < 30:
                return True
            lower = text.lower()
            return any(p in lower for p in TEMPLATE_PATTERNS)

        if len(text.strip()) < 30:
            return True

        lower = text.lower()
        if any(p in lower for p in TEMPLATE_PATTERNS):
            return True

        try:
            if detect(text) == "en":
                return True
        except Exception:
            return True

        return False

    def _extract_text_with_ocr_generator(
        self, 
        file_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None
    ) -> Generator:
        """OCR処理をジェネレータとして実行し、進捗をリアルタイムで報告"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            yield from self._extract_pdf_with_ocr_generator(file_path, file_name, ocr_engine_id, ocr_settings)
        elif ext == ".docx":
            yield from self._extract_docx_with_ocr_generator(file_path, file_name, ocr_engine_id, ocr_settings)
        elif ext == ".txt":
            yield {"file": file_name, "step": "テキストファイルを処理中..."}
            yield self._extract_text_from_txt(file_path)
        elif ext == ".csv":
            yield {"file": file_name, "step": "CSVファイルを処理中..."}
            yield self._extract_text_from_csv(file_path)
        elif ext == ".json":
            yield {"file": file_name, "step": "JSONファイルを処理中..."}
            yield self._extract_text_from_json(file_path)
        elif ext == ".eml":
            yield {"file": file_name, "step": "EMLファイルを処理中..."}
            yield self._extract_text_from_eml(file_path)
        else:
            raise ValueError(f"未対応のファイル形式: {ext}")

    def _extract_pdf_with_ocr_generator(
        self, 
        pdf_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None
    ) -> Generator:
        """PDFからテキスト抽出（ジェネレータ版、リアルタイム進捗報告）"""
        import fitz
        
        doc = fitz.open(pdf_path)
        text_list = []
        total_pages = len(doc)
        
        # OCRエンジンの決定
        if not ocr_engine_id:
            ocr_engine_id = self.ocr_factory.settings_manager.get_default_engine()
        
        yield {"file": file_name, "step": f"OCRエンジン: {ocr_engine_id} を使用"}

        for page_number in range(total_pages):
            # ページ処理開始時間を記録
            page_start_time = time.perf_counter()
            
            # 初期の「処理中...」メッセージ
            yield {
                "file": file_name, 
                "step": f"ページ {page_number + 1}/{total_pages} 処理中...",
                "page_id": f"page_{page_number + 1}_{total_pages}",  # ページ識別子
                "is_progress_update": True  # 進捗更新フラグ
            }
            
            page = doc.load_page(page_number)

            # PyMuPDF抽出（テキスト層）
            pdf_text = page.get_text().strip()

            # 定期的に経過時間を更新しながらOCR処理
            import threading
            import time as time_module
            
            # OCR処理を別スレッドで実行
            ocr_result = None
            ocr_exception = None
            
            def run_ocr():
                nonlocal ocr_result, ocr_exception
                try:
                    ocr_result = self.ocr_factory.process_with_settings(
                        engine_id=ocr_engine_id,
                        pdf_path=pdf_path,
                        page_num=page_number,
                        custom_params=ocr_settings or {}
                    )
                except Exception as e:
                    ocr_exception = e
            
            # OCR処理を開始
            ocr_thread = threading.Thread(target=run_ocr)
            ocr_thread.start()
            
            # OCR処理中は1秒ごとに経過時間を更新
            progress_count = 0
            while ocr_thread.is_alive():
                elapsed = time.perf_counter() - page_start_time
                progress_count += 1
                
                # 最初の進捗更新は表示しない（初期メッセージと重複を避ける）
                if progress_count > 1:
                    yield {
                        "file": file_name, 
                        "step": f"ページ {page_number + 1}/{total_pages} 処理中... ({elapsed:.1f}秒経過)",
                        "update_type": "progress_update",
                        "page_id": f"page_{page_number + 1}_{total_pages}",  # ページ識別子
                        "is_progress_update": True  # 進捗更新フラグ
                    }
                
                time_module.sleep(1)
                
                # 安全装置：60秒以上経過した場合は強制終了
                if elapsed > 60:
                    print(f"⚠️ ページ {page_number + 1} OCR処理が60秒を超過しました")
                    break
            
            # OCR処理完了を待機
            ocr_thread.join()
            
            # 例外が発生した場合は再発生
            if ocr_exception:
                raise ocr_exception
            
            if ocr_result["success"]:
                ocr_text = ocr_result["text"].strip()
                yield {
                    "file": file_name, 
                    "step": f"ページ {page_number + 1} OCR完了 (処理時間: {ocr_result['processing_time']:.2f}秒)"
                }
            else:
                ocr_text = f"OCRエラー: {ocr_result['error']}"
                yield {
                    "file": file_name, 
                    "step": f"ページ {page_number + 1} OCRエラー: {ocr_result['error']}"
                }

            # 構造化データの抽出
            structural_data = self._extract_structural_data(page, page_number)
            
            # タグ形式で結合して1ページ分として格納
            merged = f"<pdf_text>\n{pdf_text}\n</pdf_text>\n\n<ocr_text>\n{ocr_text}\n</ocr_text>\n\n<structural_data>\n{structural_data}\n</structural_data>"
            text_list.append(merged)

        doc.close()
        yield text_list  # 最終結果を返す

    def _extract_docx_with_ocr_generator(
        self, 
        docx_path: str, 
        file_name: str, 
        ocr_engine_id: str = None, 
        ocr_settings: Dict = None
    ) -> Generator:
        """Wordファイルから構造とOCRを統合して抽出（ジェネレータ版）"""
        import docx
        import subprocess
        import tempfile
        
        structured_text = []

        yield {"file": file_name, "step": "Word文書の構造を抽出中..."}

        # 表や段落など、論理構造ベースで抽出
        doc = docx.Document(docx_path)
        structured_text.extend([p.text for p in doc.paragraphs if p.text])

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if row_text:
                    structured_text.append(" | ".join(row_text))

        # OCRベース（LibreOffice経由でPDF化 → 新しいOCRサービス使用）
        with tempfile.TemporaryDirectory() as tmpdir:
            yield {"file": file_name, "step": "Word文書をPDFに変換中..."}
                
            try:
                subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", tmpdir, docx_path
                ], check=True)
            except subprocess.CalledProcessError:
                raise RuntimeError("LibreOfficeによるPDF変換に失敗しました")

            basename = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(tmpdir, f"{basename}.pdf")
            if os.path.exists(pdf_path):
                ocr_text = []
                for event_or_result in self._extract_pdf_with_ocr_generator(pdf_path, file_name, ocr_engine_id, ocr_settings):
                    if isinstance(event_or_result, dict) and "step" in event_or_result:
                        yield event_or_result
                    else:
                        ocr_text = event_or_result
            else:
                raise FileNotFoundError(f"PDF変換後ファイルが見つかりません: {pdf_path}")

        # 統合：タグを付けて LLM 整形に渡す
        structured_content = "\n".join(structured_text)
        ocr_content = "\n".join(ocr_text)
        
        # 両内容を統合して1つのテキストとして返す
        combined_text = f"<structured_text>\n{structured_content}\n</structured_text>\n\n<ocr_text>\n{ocr_content}\n</ocr_text>"
        yield [combined_text]

    def _run_with_timeout(self, func, timeout_sec: int):
        """タイムアウト付きで関数を実行"""
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(func)
            try:
                return future.result(timeout=timeout_sec)
            except concurrent.futures.TimeoutError:
                raise TimeoutError

    async def _apply_multimodal_processing(
        self, 
        text_content: str, 
        file_path: str, 
        file_name: str,
        progress_yield=None
    ) -> str:
        """マルチモーダル処理を適用"""
        try:
            if progress_yield:
                progress_yield({"file": file_name, "step": "マルチモーダル処理を適用中..."})
            
            # 画像ファイルの抽出
            image_paths = self._extract_images_from_pdf(file_path)
            
            # 構造情報の抽出（既存の構造化データから）
            structure_info = self._extract_structure_from_text(text_content)
            
            # マルチモーダル処理を実行
            multimodal_result = multimodal_processor.process_document(
                text_content=text_content,
                image_paths=image_paths,
                structure_info=structure_info
            )
            
            if multimodal_result["success"]:
                if progress_yield:
                    progress_yield({"file": file_name, "step": "マルチモーダル処理完了"})
                
                # 統合結果をテキストに変換
                integrated_text = self._format_multimodal_result(multimodal_result)
                return integrated_text
            else:
                LOGGER.warning(f"マルチモーダル処理エラー: {multimodal_result.get('error', 'unknown')}")
                if progress_yield:
                    progress_yield({"file": file_name, "step": f"マルチモーダル処理エラー: {multimodal_result.get('error', 'unknown')}"})
                return text_content  # 元のテキストを返す
                
        except Exception as e:
            LOGGER.exception("マルチモーダル処理例外")
            if progress_yield:
                progress_yield({"file": file_name, "step": f"マルチモーダル処理例外: {str(e)}"})
            return text_content  # 元のテキストを返す

    def _extract_images_from_pdf(self, pdf_path: str) -> List[str]:
        """PDFから画像を抽出"""
        try:
            import fitz
            import tempfile
            
            doc = fitz.open(pdf_path)
            image_paths = []
            
            # 一時ディレクトリを作成
            with tempfile.TemporaryDirectory() as tmpdir:
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # 画像リストを取得
                    image_list = page.get_images()
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            # 画像を抽出
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                            else:  # CMYK: convert to RGB
                                pix1 = fitz.Pixmap(fitz.csRGB, pix)
                                img_data = pix1.tobytes("png")
                                pix1 = None
                            
                            # 画像ファイルとして保存
                            img_filename = f"page_{page_num + 1}_img_{img_index + 1}.png"
                            img_path = os.path.join(tmpdir, img_filename)
                            
                            with open(img_path, "wb") as img_file:
                                img_file.write(img_data)
                            
                            image_paths.append(img_path)
                            pix = None
                            
                        except Exception as e:
                            LOGGER.warning(f"画像抽出エラー (page {page_num + 1}, img {img_index + 1}): {e}")
                            continue
                
                doc.close()
                return image_paths
                
        except Exception as e:
            LOGGER.exception("PDF画像抽出エラー")
            return []

    def _extract_structure_from_text(self, text_content: str) -> Dict:
        """テキストから構造情報を抽出"""
        try:
            import re
            import json
            
            structure_info = {
                "headings": [],
                "paragraphs": [],
                "lists": [],
                "tables": []
            }
            
            # <structural_data>タグからJSONデータを抽出
            structural_match = re.search(r'<structural_data>\s*(.*?)\s*</structural_data>', text_content, re.DOTALL)
            if structural_match:
                try:
                    structural_data = json.loads(structural_match.group(1))
                    
                    for item in structural_data:
                        if item.get("type") == "text_block":
                            text = item.get("text", "")
                            page = item.get("page", 1)
                            
                            # 見出し判定（#で始まる行）
                            if text.strip().startswith("#"):
                                structure_info["headings"].append({
                                    "level": text.count("#"),
                                    "text": text.strip(),
                                    "page": page,
                                    "position": item.get("bbox", [0, 0, 0, 0])[1]  # y座標
                                })
                            # リスト判定（• - 1. 2.で始まる行）
                            elif re.match(r'^[\s]*[•\-1-9\.]', text.strip()):
                                structure_info["lists"].append({
                                    "type": "bullet" if "•" in text or "-" in text else "numbered",
                                    "items": [text.strip()],
                                    "page": page,
                                    "position": item.get("bbox", [0, 0, 0, 0])[1]
                                })
                            # 段落として追加
                            else:
                                structure_info["paragraphs"].append({
                                    "text": text.strip(),
                                    "page": page,
                                    "position": item.get("bbox", [0, 0, 0, 0])[1]
                                })
                                
                except json.JSONDecodeError:
                    LOGGER.warning("構造化データのJSON解析エラー")
            
            return structure_info
            
        except Exception as e:
            LOGGER.exception("構造情報抽出エラー")
            return {"headings": [], "paragraphs": [], "lists": [], "tables": []}

    def _format_multimodal_result(self, multimodal_result: Dict) -> str:
        """マルチモーダル処理結果をテキスト形式に変換"""
        try:
            # 統合結果を構築
            integrated_parts = []
            
            # テキスト内容
            if "text_content" in multimodal_result:
                integrated_parts.append(f"<processed_text>\n{multimodal_result['text_content']}\n</processed_text>")
            
            # 画像説明
            if "image_descriptions" in multimodal_result and multimodal_result["image_descriptions"]:
                image_summary = multimodal_result.get("integrated_result", {}).get("image_summary", "画像あり")
                integrated_parts.append(f"<image_summary>\n{image_summary}\n</image_summary>")
            
            # 構造情報
            if "structure_info" in multimodal_result and multimodal_result["structure_info"]:
                structure_summary = multimodal_result.get("integrated_result", {}).get("structure_summary", "構造情報あり")
                integrated_parts.append(f"<structure_summary>\n{structure_summary}\n</structure_summary>")
            
            # マルチモーダル洞察
            if "integrated_result" in multimodal_result and "multimodal_insights" in multimodal_result["integrated_result"]:
                insights = multimodal_result["integrated_result"]["multimodal_insights"]
                if insights:
                    insights_text = "\n".join(insights)
                    integrated_parts.append(f"<multimodal_insights>\n{insights_text}\n</multimodal_insights>")
            
            return "\n\n".join(integrated_parts)
            
        except Exception as e:
            LOGGER.exception("マルチモーダル結果フォーマットエラー")
            return multimodal_result.get("text_content", "")