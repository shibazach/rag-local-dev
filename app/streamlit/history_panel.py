# app/streamlit/history_panel.py
import datetime
import streamlit as st
from io import BytesIO

def generate_history_file(format_type="txt"):
    buffer = BytesIO()
    if format_type == "txt":
        content = ""
        for idx, h in enumerate(st.session_state.history, 1):
            content += (
                f"{idx}. 質問: {h['query']}\n"
                f"モデル: {h['model']}\n"
                f"回答（{h['time']} 秒）:\n{h['response']}\n\n"
            )
        buffer.write(content.encode("utf-8"))
    elif format_type == "rtf":
        content = "{\\rtf1\\ansi\n"
        for idx, h in enumerate(st.session_state.history, 1):
            content += (
                f"\\b {idx}. 質問: {h['query']}\\b0\\line\n"
                f"モデル: {h['model']}\\line\n"
                f"\\i 回答（{h['time']} 秒）：\\i0\\line\n"
                f"{h['response']}\\line\\line\n"
            )
        content += "}"
        buffer.write(content.encode("utf-8"))
    else:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for idx, h in enumerate(st.session_state.history, 1):
            doc.add_heading(f"{idx}. 質問: {h['query']}", level=2)
            doc.add_paragraph(f"モデル: {h['model']}")
            p = doc.add_paragraph()
            run = p.add_run(f"回答（{h['time']} 秒）:\n{h['response']}")
            run.font.size = Pt(11)
        doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_history_panel():
    st.markdown("### 📜 検索履歴")
    if st.session_state.history:
        for i, item in enumerate(reversed(st.session_state.history), 1):
            with st.expander(f"{i}. 質問: {item['query']}（モデル: {item['model']}）"):
                st.markdown(f"🧠 回答（{item['time']} 秒）:\n\n{item['response']}")
    else:
        st.write("まだ履歴はありません。")

    st.markdown("### 💾 履歴をファイルに保存")
    format_type = st.selectbox("保存形式を選択してください", ["txt", "rtf", "docx"])
    if st.button("📥 ダウンロード"):
        file_buffer = generate_history_file(format_type)
        mime_map = {
            "txt": "text/plain",
            "rtf": "application/rtf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        st.download_button(
            "ファイルをダウンロード",
            data=file_buffer,
            file_name=f"search_history_{today_str}.{format_type}",
            mime=mime_map[format_type],
        )
