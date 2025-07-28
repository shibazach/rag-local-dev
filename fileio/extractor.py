# fileio/extractor.py  最終更新 2025-07-08 00:00
import os, json, csv, subprocess, tempfile
from io import BytesIO
from typing import List

import docx
import fitz  # PyMuPDF
import numpy as np
import pytesseract
from PIL import Image

from src import bootstrap
from ocr import (correct_orientation, detect_rotation_angle)

# PDFを画像に変換（ページごと）
def pdf_page_to_image(pdf_path, page_number, dpi=300) -> Image.Image:
    doc = fitz.open(pdf_path)
    page = doc.load_page(page_number)
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat)
    img_data = pix.tobytes("png")
    return Image.open(BytesIO(img_data))

# PDFからテキスト（PyMuPDF + OCR 両方、OCRは向き補正つき）
def extract_text_from_pdf(pdf_path: str, progress_callback=None) -> List[str]:
    doc = fitz.open(pdf_path)
    text_list = []
    total_pages = len(doc)

    for page_number in range(total_pages):
        # 進捗コールバック
        if progress_callback:
            progress_callback(f"ページ {page_number + 1}/{total_pages} 処理中...")

        page = doc.load_page(page_number)

        # PyMuPDF抽出（テキスト層）
        pdf_text = page.get_text().strip()

        # OCR抽出（画像層 + 向き補正）
        image = pdf_page_to_image(pdf_path, page_number)
        angle = detect_rotation_angle(image)
        rotated = correct_orientation(image, angle)
        ocr_text = pytesseract.image_to_string(rotated, lang="jpn+eng").strip()

        # 結合して1ページ分として格納
        merged = f"【PDF抽出】\n{pdf_text}\n\n【OCR抽出】\n{ocr_text}"
        text_list.append(merged)

    return text_list

# PDFからテキスト（ジェネレータ版 - 進捗をyield）
def extract_text_from_pdf_with_progress(pdf_path: str):
    """PDFテキスト抽出をページごとに進捗をyieldしながら実行"""
    doc = fitz.open(pdf_path)
    text_list = []
    total_pages = len(doc)

    for page_number in range(total_pages):
        # 進捗をyield
        yield {"progress": f"ページ {page_number + 1}/{total_pages} 処理中..."}

        page = doc.load_page(page_number)

        # PyMuPDF抽出（テキスト層）
        pdf_text = page.get_text().strip()

        # OCR抽出（画像層 + 向き補正）
        image = pdf_page_to_image(pdf_path, page_number)
        angle = detect_rotation_angle(image)
        rotated = correct_orientation(image, angle)
        ocr_text = pytesseract.image_to_string(rotated, lang="jpn+eng").strip()

        # 結合して1ページ分として格納
        merged = f"【PDF抽出】\n{pdf_text}\n\n【OCR抽出】\n{ocr_text}"
        text_list.append(merged)

    # 最終結果をyield
    yield {"result": text_list}

# Wordファイルから構造とOCRを統合して抽出
def extract_text_from_docx_combined(docx_path: str) -> List[str]:
    structured_text = []

    # 🔹 表や段落など、論理構造ベースで抽出
    doc = docx.Document(docx_path)
    structured_text.extend([p.text for p in doc.paragraphs if p.text])

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
            if row_text:
                structured_text.append(" | ".join(row_text))

    # 🔹 OCRベース（LibreOffice経由でPDF化 → extract_text_from_pdf を流用）
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, docx_path
            ], check=True)
        except subprocess.CalledProcessError:
            raise RuntimeError("LibreOfficeによるPDF変換に失敗しました")

        basename = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(tmpdir, f"{basename}.pdf")
        if os.path.exists(pdf_path):
            ocr_text = extract_text_from_pdf(pdf_path)
        else:
            raise FileNotFoundError(f"PDF変換後ファイルが見つかりません: {pdf_path}")

    # 🔹 統合：タグを付けて LLM 整形に渡す
    structured_content = "\n".join(structured_text)
    ocr_content = "\n".join(ocr_text)
    
    # 両内容を統合して1つのテキストとして返す
    combined_text = f"<structured_text>\n{structured_content}\n</structured_text>\n\n<ocr_text>\n{ocr_content}\n</ocr_text>"
    return [combined_text]

# TXT
def extract_text_from_txt(txt_path: str) -> List[str]:
    with open(txt_path, encoding="utf-8") as f:
        return [line.strip() for line in f if line.strip()]

# CSV（textカラムがなければログに記録してスキップ）
def extract_text_from_csv(csv_path: str) -> List[str]:
    results = []
    try:
        with open(csv_path, encoding="utf-8") as f:
            reader = csv.DictReader(f)
            if "text" not in reader.fieldnames:
                raise ValueError("CSVに'text'カラムが必要です")
            for row in reader:
                if row.get("text"):
                    results.append(row["text"].strip())
        return results
    except Exception as e:
        os.makedirs("logs", exist_ok=True)
        log_path = os.path.join("logs", "invalid_csv_log.txt")
        from datetime import datetime
        now = datetime.utcnow().isoformat() + "Z"
        with open(log_path, "a", encoding="utf-8") as logf:
            logf.write(f"{now}: Skipped {csv_path} - {e}\n")
        return []

# EML（プレーンテキスト部分のみを対象、引用除去）
def extract_text_from_eml(file_path: str) -> List[str]:
    from email import policy
    from email.parser import BytesParser

    with open(file_path, "rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)

    text_blocks = []
    for part in msg.walk():
        if part.get_content_type() == "text/plain":
            content = part.get_content()
            lines = content.splitlines()
            # REM: 引用行（>>で始まる行）は除去
            lines = [line for line in lines if not line.strip().startswith(">>")]
            filtered = "\n".join(lines).strip()
            if filtered:
                text_blocks.append(filtered)
    return text_blocks

# JSON
def extract_text_from_json(json_path: str) -> List[str]:
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
        if not isinstance(data, list):
            raise ValueError("JSONはリスト形式である必要があります")
        return [item["text"] for item in data if "text" in item]

# 拡張子で振り分け
def extract_text_by_extension(file_path: str, progress_callback=None) -> List[str]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path, progress_callback)
    elif ext == ".docx":
        if progress_callback:
            progress_callback("Word文書を処理中...")
        return extract_text_from_docx_combined(file_path)
    elif ext == ".txt":
        if progress_callback:
            progress_callback("テキストファイルを処理中...")
        return extract_text_from_txt(file_path)
    elif ext == ".csv":
        if progress_callback:
            progress_callback("CSVファイルを処理中...")
        return extract_text_from_csv(file_path)
    elif ext == ".json":
        if progress_callback:
            progress_callback("JSONファイルを処理中...")
        return extract_text_from_json(file_path)
    else:
        raise ValueError(f"未対応のファイル形式: {ext}")
