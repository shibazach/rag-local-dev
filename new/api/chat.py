# REM: api/chat.py @2024-12-19
# REM: チャット・検索API（関数型実装、OLD系完全移植）

# ── 標準ライブラリ ──
import os
import json
import asyncio
import datetime
from typing import Dict, List, Optional
from io import BytesIO

# ── サードパーティ ──
from fastapi import APIRouter, Form, HTTPException, Depends, Request
from fastapi.responses import JSONResponse, StreamingResponse

# ── プロジェクト内 ──
from new.config import LOGGER
from new.schemas import SuccessResponse, ErrorResponse
from new.auth_functions import require_authentication

# ──────────────────────────────────────────────────────────
# ルーター設定
# ──────────────────────────────────────────────────────────
router = APIRouter(prefix="/chat", tags=["chat"])

# ──────────────────────────────────────────────────────────
# グローバル変数（検索処理制御）
# ──────────────────────────────────────────────────────────
_current_search_task = None
_search_cancelled = False

# ──────────────────────────────────────────────────────────
# 履歴管理関数群（Pure Functions）
# ──────────────────────────────────────────────────────────

def get_history_file_path() -> str:
    """履歴ファイルパスを取得"""
    return "logs/search_history.json"

async def save_search_history(
    query: str, 
    model_key: str, 
    mode: str, 
    result: Dict, 
    processing_time: float
) -> None:
    """
    検索履歴を保存（非同期、原子的書き込み）
    
    Args:
        query: 検索クエリ
        model_key: 使用した埋め込みモデル
        mode: 検索モード
        result: 検索結果
        processing_time: 処理時間（秒）
    """
    try:
        history_file = get_history_file_path()
        
        # 履歴ディレクトリを作成
        os.makedirs(os.path.dirname(history_file), exist_ok=True)
        
        # 既存の履歴を読み込み
        history = []
        if os.path.exists(history_file):
            try:
                with open(history_file, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if content:
                        history = json.loads(content)
                    else:
                        history = []
            except (json.JSONDecodeError, UnicodeDecodeError) as e:
                LOGGER.warning(f"履歴ファイル読み込みエラー（破損ファイルを初期化）: {e}")
                history = []
        
        # 新しい履歴エントリを作成
        history_entry = {
            "id": len(history) + 1,
            "timestamp": datetime.datetime.now().isoformat(),
            "query": query,
            "model_key": model_key,
            "mode": mode,
            "processing_time": processing_time,
            "result": result
        }
        
        # 履歴に追加（最新を先頭に）
        history.insert(0, history_entry)
        
        # 履歴を最大100件に制限
        if len(history) > 100:
            history = history[:100]
        
        # ファイルに保存（原子的書き込みで破損を防止）
        import tempfile
        temp_file = history_file + '.tmp'
        try:
            with open(temp_file, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
            # 原子的にファイルを置き換え
            os.replace(temp_file, history_file)
        except Exception as e:
            # 一時ファイルが残っている場合は削除
            if os.path.exists(temp_file):
                os.remove(temp_file)
            raise e
            
    except Exception as e:
        LOGGER.error(f"履歴保存エラー: {e}")

async def load_search_history() -> List[Dict]:
    """
    検索履歴を読み込み
    
    Returns:
        List[Dict]: 検索履歴リスト
    """
    try:
        history_file = get_history_file_path()
        
        if not os.path.exists(history_file):
            return []
        
        with open(history_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        LOGGER.error(f"履歴読み込みエラー: {e}")
        return []

async def remove_search_history_item(history_id: int) -> None:
    """
    指定された履歴を削除
    
    Args:
        history_id: 削除する履歴のID
    """
    try:
        history = await load_search_history()
        history = [h for h in history if h.get("id") != history_id]
        
        history_file = get_history_file_path()
        with open(history_file, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
    except Exception as e:
        LOGGER.error(f"履歴削除エラー: {e}")

async def clear_all_search_history() -> None:
    """全履歴を削除"""
    try:
        history_file = get_history_file_path()
        if os.path.exists(history_file):
            os.remove(history_file)
    except Exception as e:
        LOGGER.error(f"履歴全削除エラー: {e}")

# ──────────────────────────────────────────────────────────
# 検索処理関数群
# ──────────────────────────────────────────────────────────

async def handle_search_query(
    query: str,
    model_key: str,
    mode: str = "チャンク統合",
    search_limit: int = 10,
    min_score: float = 0.0
) -> Dict:
    """
    検索クエリを処理（スタブ実装）
    
    Args:
        query: 検索クエリ
        model_key: 埋め込みモデルキー
        mode: 検索モード
        search_limit: 検索件数上限
        min_score: 最小一致度
    
    Returns:
        Dict: 検索結果
    """
    # TODO: 実際の検索処理をOLD系から移植
    # 現在はスタブ実装
    
    await asyncio.sleep(1)  # 検索処理をシミュレート
    
    if mode == "チャンク統合":
        return {
            "mode": mode,
            "answer": f"「{query}」に関する統合回答です。（スタブ実装）",
            "sources": [
                {"file_id": "sample1", "file_name": "サンプル文書1.pdf"},
                {"file_id": "sample2", "file_name": "サンプル文書2.pdf"}
            ],
            "results": []
        }
    else:
        return {
            "mode": mode,
            "summaries": [
                {
                    "file_id": "sample1",
                    "file_name": "サンプル文書1.pdf",
                    "score": 0.85,
                    "summary": f"「{query}」に関連する内容のサマリー1"
                },
                {
                    "file_id": "sample2", 
                    "file_name": "サンプル文書2.pdf",
                    "score": 0.72,
                    "summary": f"「{query}」に関連する内容のサマリー2"
                }
            ],
            "results": []
        }

# ──────────────────────────────────────────────────────────
# ファイル生成関数群
# ──────────────────────────────────────────────────────────

async def generate_history_file(history: List[Dict], format_type: str) -> bytes:
    """
    履歴ファイルを生成
    
    Args:
        history: 履歴データ
        format_type: ファイル形式（txt, json, rtf, docx）
    
    Returns:
        bytes: 生成されたファイルの内容
    """
    if format_type == "txt":
        content = ""
        for idx, h in enumerate(history, 1):
            timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
            content += (
                f"{idx}. 質問: {h['query']}\n"
                f"モデル: {h['model_key']} | モード: {h['mode']}\n"
                f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}\n"
            )
            
            # 結果の表示
            result = h.get('result', {})
            if result.get('mode') == 'チャンク統合':
                content += f"統合回答:\n{result.get('answer', '')}\n"
                if result.get('sources'):
                    content += "参照ファイル:\n"
                    for src in result['sources']:
                        content += f"  - {src.get('file_name', '')}\n"
            else:
                content += "ファイル別結果:\n"
                for res in result.get('summaries', []):
                    content += f"  - {res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\n"
                    content += f"    要約: {res.get('summary', '')}\n"
            content += "\n" + "="*50 + "\n\n"
        
        return content.encode("utf-8")
    
    elif format_type == "json":
        return json.dumps(history, ensure_ascii=False, indent=2).encode("utf-8")
    
    elif format_type == "rtf":
        content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24\n"
        for idx, h in enumerate(history, 1):
            timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
            content += (
                f"\\b {idx}. 質問: {h['query']}\\b0\\line\n"
                f"モデル: {h['model_key']} | モード: {h['mode']}\\line\n"
                f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}\\line\\line\n"
            )
            
            result = h.get('result', {})
            if result.get('mode') == 'チャンク統合':
                content += f"\\i 統合回答:\\i0\\line\n{result.get('answer', '')}\\line\\line\n"
            else:
                content += "\\i ファイル別結果:\\i0\\line\n"
                for res in result.get('summaries', []):
                    content += f"{res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\\line\n"
            content += "\\line\n"
        content += "}"
        return content.encode("utf-8")
    
    elif format_type == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            doc.add_heading('検索履歴', 0)
            
            for idx, h in enumerate(history, 1):
                timestamp = datetime.datetime.fromisoformat(h['timestamp']).strftime("%Y-%m-%d %H:%M:%S")
                
                # 質問
                doc.add_heading(f"{idx}. 質問: {h['query']}", level=2)
                
                # メタ情報
                meta_p = doc.add_paragraph()
                meta_p.add_run(f"モデル: {h['model_key']} | モード: {h['mode']}\n")
                meta_p.add_run(f"処理時間: {h['processing_time']}秒 | 実行日時: {timestamp}")
                
                # 結果
                result = h.get('result', {})
                if result.get('mode') == 'チャンク統合':
                    doc.add_heading('統合回答', level=3)
                    answer_p = doc.add_paragraph(result.get('answer', ''))
                    answer_p.style.font.size = Pt(11)
                else:
                    doc.add_heading('ファイル別結果', level=3)
                    for res in result.get('summaries', []):
                        file_p = doc.add_paragraph()
                        file_p.add_run(f"{res.get('file_name', '')} (一致度: {res.get('score', 0):.2f})\n").bold = True
                        file_p.add_run(f"要約: {res.get('summary', '')}")
                
                doc.add_page_break()
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except ImportError:
            # python-docxがインストールされていない場合はテキスト形式で返す
            return await generate_history_file(history, "txt")
    
    else:
        raise ValueError(f"サポートされていない形式: {format_type}")

# ──────────────────────────────────────────────────────────
# API エンドポイント群
# ──────────────────────────────────────────────────────────

@router.post("/query")
async def search_query_endpoint(
    request: Request,
    query: str = Form(...),
    model_key: str = Form(...),
    mode: str = Form("チャンク統合"),
    search_limit: int = Form(10),
    min_score: float = Form(0.0),
    current_user = Depends(require_authentication)
) -> JSONResponse:
    """検索クエリ実行（関数型実装）"""
    global _current_search_task, _search_cancelled
    
    # ガードクローズ: 入力検証
    if not query.strip():
        raise HTTPException(400, "検索クエリが空です")
    
    if search_limit < 1 or search_limit > 50:
        raise HTTPException(400, "検索件数は1-50の範囲で指定してください")
    
    try:
        import time
        
        # キャンセルフラグをリセット
        _search_cancelled = False
        
        start_time = time.time()
        
        # キャンセル可能なタスクとして実行
        _current_search_task = asyncio.create_task(
            handle_search_query(query, model_key, mode, search_limit, min_score)
        )
        
        result = await _current_search_task
        processing_time = round(time.time() - start_time, 2)
        
        # 履歴に保存
        await save_search_history(query, model_key, mode, result, processing_time)
        
        return JSONResponse(result)
        
    except asyncio.CancelledError:
        return JSONResponse({"error": "検索処理がキャンセルされました"}, status_code=499)
    except Exception as e:
        LOGGER.exception(f"検索処理エラー: {e}")
        raise HTTPException(500, f"検索処理でエラーが発生しました: {str(e)}")
    finally:
        _current_search_task = None
        _search_cancelled = False

@router.get("/history")
async def get_search_history_endpoint(
    request: Request,
    current_user = Depends(require_authentication)
) -> List[Dict]:
    """検索履歴取得"""
    try:
        return await load_search_history()
    except Exception as e:
        LOGGER.exception(f"履歴取得エラー: {e}")
        raise HTTPException(500, f"履歴取得でエラーが発生しました: {str(e)}")

@router.delete("/history/{history_id}")
async def delete_search_history_endpoint(
    history_id: int,
    request: Request,
    current_user = Depends(require_authentication)
) -> SuccessResponse:
    """指定された履歴を削除"""
    try:
        await remove_search_history_item(history_id)
        return SuccessResponse(message="履歴を削除しました")
    except Exception as e:
        LOGGER.exception(f"履歴削除エラー: {e}")
        raise HTTPException(500, f"履歴削除でエラーが発生しました: {str(e)}")

@router.delete("/history")
async def clear_search_history_endpoint(
    request: Request,
    current_user = Depends(require_authentication)
) -> SuccessResponse:
    """全履歴を削除"""
    try:
        await clear_all_search_history()
        return SuccessResponse(message="全履歴を削除しました")
    except Exception as e:
        LOGGER.exception(f"履歴全削除エラー: {e}")
        raise HTTPException(500, f"履歴全削除でエラーが発生しました: {str(e)}")

@router.post("/stop_search")
async def stop_current_search_endpoint(
    request: Request,
    current_user = Depends(require_authentication)
) -> SuccessResponse:
    """現在実行中の検索処理を停止"""
    global _search_cancelled
    
    try:
        _search_cancelled = True
        
        # 現在のタスクがあればキャンセル
        if _current_search_task and not _current_search_task.done():
            _current_search_task.cancel()
            
        return SuccessResponse(message="検索処理の停止を要求しました")
    except Exception as e:
        LOGGER.exception(f"検索停止エラー: {e}")
        raise HTTPException(500, f"検索停止でエラーが発生しました: {str(e)}")

@router.get("/history/download/{format_type}")
async def download_search_history_endpoint(
    format_type: str,
    request: Request,
    current_user = Depends(require_authentication)
) -> StreamingResponse:
    """検索履歴をダウンロード"""
    # ガードクローズ: フォーマット検証
    if format_type not in ["txt", "json", "rtf", "docx"]:
        raise HTTPException(400, f"サポートされていない形式: {format_type}")
    
    try:
        history = await load_search_history()
        if not history:
            raise HTTPException(404, "履歴がありません")
        
        # ファイル生成
        buffer = await generate_history_file(history, format_type)
        
        # MIMEタイプとファイル名を設定
        mime_map = {
            "txt": "text/plain",
            "rtf": "application/rtf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "json": "application/json"
        }
        
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        filename = f"search_history_{today_str}.{format_type}"
        
        return StreamingResponse(
            BytesIO(buffer),
            media_type=mime_map.get(format_type, "application/octet-stream"),
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        LOGGER.exception(f"履歴ダウンロードエラー: {e}")
        raise HTTPException(500, f"履歴ダウンロードでエラーが発生しました: {str(e)}")