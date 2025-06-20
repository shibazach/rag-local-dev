# scripts/rag_chat.py

import time
import datetime
import threading
import streamlit as st
import numpy as np
from sqlalchemy import create_engine, text
from langchain_community.embeddings import OllamaEmbeddings
from langchain_ollama import OllamaLLM
from sentence_transformers import SentenceTransformer
from io import BytesIO

from src import bootstrap
from src.embedding_config import embedding_options
from src.error_handler import install_global_exception_handler
from src.embedder import embed_and_insert

install_global_exception_handler()

# ğŸ†• Ollamaã‚’ã‚¹ãƒªãƒ¼ãƒ—ã•ã›ãªã„ãƒãƒƒã‚¯ã‚°ãƒ©ã‚¦ãƒ³ãƒ‰ã‚¹ãƒ¬ãƒƒãƒ‰
def keep_ollama_warm():
    llm_ping = OllamaLLM(model="phi4-mini", base_url="http://172.18.0.1:11434")
    while True:
        try:
            llm_ping.invoke("ping")
            print("âœ… Ollama keep-alive sent.")
        except Exception as e:
            print("âš ï¸ Ollama keep-alive failed:", e)
        time.sleep(600)  # 10åˆ†ã”ã¨ã«Ping

# ğŸ” èµ·å‹•æ™‚ã«é–‹å§‹ï¼ˆdaemon=True ãªã®ã§ãƒãƒƒã‚¯ã‚°ãƒ©ã‚¦ãƒ³ãƒ‰ï¼‰
threading.Thread(target=keep_ollama_warm, daemon=True).start()

st.set_page_config(page_title="ãƒ­ãƒ¼ã‚«ãƒ«RAGãƒãƒ£ãƒƒãƒˆ", layout="wide")

engine = create_engine("postgresql://raguser:ragpass@pgvector-db:5432/ragdb")

if "history" not in st.session_state:
    st.session_state.history = []
if "query_input" not in st.session_state:
    st.session_state.query_input = ""

st.title("ğŸ” ãƒ­ãƒ¼ã‚«ãƒ«RAGãƒãƒ£ãƒƒãƒˆ")

model_choices = {
    f"{v['model_name']} ({v['embedder']})": k
    for k, v in embedding_options.items()
}
selected_label = st.selectbox("ä½¿ç”¨ã™ã‚‹åŸ‹ã‚è¾¼ã¿ãƒ¢ãƒ‡ãƒ«ã‚’é¸æŠã—ã¦ãã ã•ã„ï¼š", list(model_choices.keys()))
selected_key = model_choices[selected_label]
selected_model = embedding_options[selected_key]

model_safe = selected_model["model_name"].replace("/", "_").replace("-", "_")
tablename = f"{model_safe}_{selected_model['dimension']}"

def to_pgvector_literal(vec):
    if isinstance(vec, np.ndarray):
        vec = vec.tolist()
    return "[" + ",".join(f"{float(x):.6f}" for x in vec) + "]"

def search_similar_documents(query_embedding, top_k=5):
    embedding_str = to_pgvector_literal(query_embedding)
    sql = f"""
        SELECT
          e.content   AS snippet,
          e.file_id   AS file_id,
          f.filename  AS filename,
          f.content   AS full_text,
          f.file_blob AS file_blob
        FROM "{tablename}" AS e
        JOIN files AS f
          ON e.file_id = f.file_id
        ORDER BY e.embedding <-> '{embedding_str}'::vector
        LIMIT :top_k
    """
    with engine.connect() as conn:
        result = conn.execute(text(sql), {"top_k": top_k})
        return [dict(row) for row in result]

llm = OllamaLLM(model="phi4-mini", base_url="http://172.18.0.1:11434")

left_col, right_col = st.columns([1, 1])

with left_col:
    st.markdown("#### ğŸ” è³ªå•ã‚’å…¥åŠ›ã—ã¦æ¤œç´¢ã‚’å®Ÿè¡Œ")
    st.session_state.query_input = st.text_area("è³ªå•ã‚’å…¥åŠ›ã—ã¦ãã ã•ã„:", value=st.session_state.query_input, height=100)

    if st.button("æ¤œç´¢å®Ÿè¡Œ"):
        query = st.session_state.query_input.strip()
        if query:
            start_time = time.time()

            if selected_model["embedder"] == "OllamaEmbeddings":
                embedder = OllamaEmbeddings(model=selected_model["model_name"], base_url="http://172.18.0.1:11434")
                query_embedding = embedder.embed_query(query)
            else:
                embedder = SentenceTransformer(selected_model["model_name"])
                query_embedding = embedder.encode([query], convert_to_numpy=True)[0]

            docs = search_similar_documents(query_embedding)

            context = "\n\n".join(f"\U0001F4C4 **{d['filename']}**: {d['snippet']}" for d in docs)

            prompt = f"""
ä»¥ä¸‹ã®æƒ…å ±ã‚’å‚è€ƒã«ã—ã¦ã€è³ªå•ã«ç­”ãˆã¦ãã ã•ã„ã€‚

æƒ…å ±:
{context}

è³ªå•:
{query}
""".strip()

            with st.spinner("è€ƒãˆä¸­..."):
                response = llm.invoke(prompt)

            elapsed_time = round(time.time() - start_time, 2)

            st.markdown(f"### \U0001F9E0 å›ç­”ï¼ˆ{elapsed_time} ç§’ï¼‰")
            st.write(response)

            st.session_state.history.append({
                "query": query,
                "response": response,
                "model": selected_label,
                "time": elapsed_time
            })

            st.markdown("### \U0001F4D1 æ¤œç´¢çµæœãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼")
            for d in docs:
                st.markdown(f"**{d['filename']}**")
                st.write(d["snippet"])
                with st.expander("â–¶ï¸ å…¨æ–‡ã‚’ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼"):
                    edited_text = st.text_area("å…¨æ–‡ãƒ†ã‚­ã‚¹ãƒˆï¼ˆç·¨é›†å¯èƒ½ï¼‰", value=d["full_text"], height=200, key=f"edit_{d['file_id']}")

                    if st.button("ä¿å­˜ã—ã¦å†ãƒ™ã‚¯ãƒˆãƒ«åŒ–", key=f"save_{d['file_id']}"):
                        try:
                            with engine.begin() as conn:
                                conn.execute(text("UPDATE files SET content = :content WHERE file_id = :file_id"),
                                             {"content": edited_text, "file_id": d["file_id"]})
                                conn.execute(text(f'DELETE FROM "{tablename}" WHERE file_id = :file_id'),
                                             {"file_id": d["file_id"]})
                            st.success("âœ… contentã‚’æ›´æ–°ã—ã€æ—§ãƒ™ã‚¯ãƒˆãƒ«ã‚’å‰Šé™¤ã—ã¾ã—ãŸ")

                            embed_and_insert(texts=[edited_text], filename=d["filename"], truncate_done_tables=set())
                            st.success("âœ… å†ãƒ™ã‚¯ãƒˆãƒ«åŒ–å®Œäº†ï¼")
                        except Exception as e:
                            st.error(f"âŒ ã‚¨ãƒ©ãƒ¼ãŒç™ºç”Ÿã—ã¾ã—ãŸ: {e}")

                    st.download_button("å…ƒãƒ•ã‚¡ã‚¤ãƒ«ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰", data=d["file_blob"], file_name=d["filename"], mime="application/octet-stream")

with right_col:
    st.markdown("### \U0001F4DC æ¤œç´¢å±¥æ­´")
    if st.session_state.history:
        for i, item in enumerate(reversed(st.session_state.history), 1):
            with st.expander(f"{i}. è³ªå•: {item['query']}ï¼ˆãƒ¢ãƒ‡ãƒ«: {item['model']}ï¼‰", expanded=False):
                st.markdown(f"\U0001F9E0 å›ç­”ï¼ˆ{item['time']} ç§’ï¼‰:\n\n{item['response']}")
    else:
        st.write("ã¾ã å±¥æ­´ã¯ã‚ã‚Šã¾ã›ã‚“ã€‚")

    st.markdown("### \U0001F4BE å±¥æ­´ã‚’ãƒ•ã‚¡ã‚¤ãƒ«ã«ä¿å­˜")
    format_type = st.selectbox("ä¿å­˜å½¢å¼ã‚’é¸æŠã—ã¦ãã ã•ã„", ["txt", "rtf", "docx"])
    if st.button("ğŸ“¥ ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰"):
        def generate_history_file(format_type="txt"):
            buffer = BytesIO()
            if format_type == "txt":
                content = ""
                for idx, h in enumerate(st.session_state.history, 1):
                    content += f"{idx}. è³ªå•: {h['query']}\nãƒ¢ãƒ‡ãƒ«: {h['model']}\nå›ç­”ï¼ˆ{h['time']} ç§’ï¼‰:\n{h['response']}\n\n"
                buffer.write(content.encode("utf-8"))
            elif format_type == "rtf":
                content = "{\\rtf1\\ansi\n"
                for idx, h in enumerate(st.session_state.history, 1):
                    content += f"\\b {idx}. è³ªå•: {h['query']}\\b0\\line\nãƒ¢ãƒ‡ãƒ«: {h['model']}\\line\n\\i å›ç­”ï¼ˆ{h['time']} ç§’ï¼‰:\\i0\\line\n{h['response']}\\line\\line\n"
                content += "}"
                buffer.write(content.encode("utf-8"))
            else:
                from docx import Document
                from docx.shared import Pt
                doc = Document()
                for idx, h in enumerate(st.session_state.history, 1):
                    doc.add_heading(f"{idx}. è³ªå•: {h['query']}", level=2)
                    doc.add_paragraph(f"ãƒ¢ãƒ‡ãƒ«: {h['model']}")
                    p = doc.add_paragraph()
                    run = p.add_run(f"å›ç­”ï¼ˆ{h['time']} ç§’ï¼‰:\n{h['response']}")
                    run.font.size = Pt(11)
                doc.save(buffer)
            buffer.seek(0)
            return buffer

        file_buffer = generate_history_file(format_type)
        mime_map = {
            "txt": "text/plain",
            "rtf": "application/rtf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        today_str = datetime.date.today().strftime("%Y-%m-%d")
        st.download_button("ãƒ•ã‚¡ã‚¤ãƒ«ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰", data=file_buffer, file_name=f"search_history_{today_str}.{format_type}", mime=mime_map[format_type])
