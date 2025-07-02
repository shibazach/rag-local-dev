# chat_view.py（最新版：処理フロー・補助関数反映、既存機能＋ファイル別モード対応）

import streamlit as st
import time, datetime, threading
import numpy as np
import uuid, os
from sqlalchemy import create_engine, text
from io import BytesIO
from langchain_community.embeddings import OllamaEmbeddings
from langchain_ollama import OllamaLLM
from sentence_transformers import SentenceTransformer

from src import bootstrap
from src.embedder import embed_and_insert
from src.config import DB_ENGINE as engine
from src.error_handler import install_global_exception_handler
from src.config import OLLAMA_MODEL, OLLAMA_BASE, EMBEDDING_OPTIONS

install_global_exception_handler()

LLM = OllamaLLM(model=OLLAMA_MODEL, base_url=OLLAMA_BASE)

# REM: Ollamaを維持するためのping送信スレッド
def keep_ollama_warm():
    while True:
        try:
            LLM.invoke("ping")
            print("✅ Ollama keep-alive sent.")
        except Exception as e:
            print("⚠️ Ollama keep-alive failed:", e)
        time.sleep(600)

threading.Thread(target=keep_ollama_warm, daemon=True).start()

# REM: numpy配列をPostgreSQLのpgvector形式に変換
def to_pgvector_literal(vec):
    if isinstance(vec, np.ndarray):
        vec = vec.tolist()
    return "[" + ",".join(f"{float(x):.6f}" for x in vec) + "]"

# REM: UI部品用の一意なキーを生成
def make_unique_key(file_id, filename):
    safe = filename.replace(".", "_").replace(" ", "_")
    return f"{file_id}_{safe}_{uuid.uuid4().hex[:6]}"

# REM: 検索用の類似チャンク抽出（チャンク統合モード）
def search_similar_documents(query_embedding, tablename, top_k=5):
    embedding_str = to_pgvector_literal(query_embedding)
    sql = f"""
        SELECT e.content AS snippet, e.file_id, f.filename, f.content AS full_text, f.file_blob
        FROM "{tablename}" AS e
        JOIN files AS f ON e.file_id = f.file_id
        ORDER BY e.embedding <-> '{embedding_str}'::vector
        LIMIT :top_k
    """
    with engine.connect() as conn:
        result = conn.execute(text(sql), {"top_k": top_k})
        return [dict(row) for row in result.mappings()]

# REM: ファイル全文取得（file_id指定）
def get_file_content(file_id):
    with engine.connect() as conn:
        result = conn.execute(text("SELECT content FROM files WHERE file_id = :fid"), {"fid": file_id})
        row = result.fetchone()
        return row[0] if row else ""

# REM: 質問と文書の一致度スコアと要約を取得
def llm_summarize_with_score(query, content):
    prompt = f"""
以下はユーザーの質問と文書の内容です。

【質問】
{query}

【文書】
{content[:3000]}

この文書が質問にどの程度一致しているか（0.0〜1.0）で評価し、次の形式で出力してください：

一致度: <score>
要約: <summary>
"""
    result = LLM.invoke(prompt)
    import re
    m = re.search(r"一致度[:：]\s*([0-9.]+).*?要約[:：]\s*(.+)", result, re.DOTALL)
    if m:
        return m.group(2).strip(), float(m.group(1))
    return result.strip(), 0.0

# REM: 検索履歴のファイル生成（txt/rtf/docx）
def generate_history_file(format_type="txt"):
    buffer = BytesIO()
    if format_type == "txt":
        content = ""
        for idx, h in enumerate(st.session_state.history, 1):
            content += f"{idx}. 質問: {h['query']}\nモデル: {h['model']}\n回答（{h['time']} 秒）:\n{h['response']}\n\n"
        buffer.write(content.encode("utf-8"))
    elif format_type == "rtf":
        content = "{\\rtf1\\ansi\n"
        for idx, h in enumerate(st.session_state.history, 1):
            content += f"\\b {idx}. 質問: {h['query']}\\b0\\line\nモデル: {h['model']}\\line\n\\i 回答（{h['time']} 秒）：\\i0\\line\n{h['response']}\\line\\line\n"
        content += "}"
        buffer.write(content.encode("utf-8"))
    else:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for idx, h in enumerate(st.session_state.history, 1):
            doc.add_heading(f"{idx}. 質問: {h['query']}", level=2)
            doc.add_paragraph(f"モデル: {h['model']}")
            p = doc.add_paragraph()
            run = p.add_run(f"回答（{h['time']} 秒）:\n{h['response']}")
            run.font.size = Pt(11)
        doc.save(buffer)
    buffer.seek(0)
    return buffer

# REM: メイン画面の描画

def render_chat_view():
    st.title("\U0001F50D ローカルRAGチャット")

    if "history" not in st.session_state:
        st.session_state.history = []
    if "query_input" not in st.session_state:
        st.session_state.query_input = ""

    rag_mode = st.radio("検索モードを選択：", ("チャンク統合", "ファイル別（要約+一致度）"), horizontal=True)
    model_choices = {f"{v['model_name']} ({v['embedder']})": k for k, v in EMBEDDING_OPTIONS.items()}
    selected_label = st.selectbox("使用する埋め込みモデルを選択してください：", list(model_choices.keys()))
    selected_key = model_choices[selected_label]
    selected_model = EMBEDDING_OPTIONS[selected_key]
    model_safe = selected_model["model_name"].replace("/", "_").replace("-", "_")
    tablename = f"{model_safe}_{selected_model['dimension']}"

    left_col, right_col = st.columns([1, 1])

    with left_col:
        st.markdown("#### \U0001F50E 質問を入力して検索を実行")
        st.session_state.query_input = st.text_area("質問を入力してください:", value=st.session_state.query_input, height=100)

        if st.button("検索実行"):
            query = st.session_state.query_input.strip()
            if query:
                start_time = time.time()

                if selected_model["embedder"] == "OllamaEmbeddings":
                    embedder = OllamaEmbeddings(model=selected_model["model_name"], base_url=OLLAMA_BASE)
                    query_embedding = embedder.embed_query(query)
                else:
                    embedder = SentenceTransformer(selected_model["model_name"])
                    query_embedding = embedder.encode([query], convert_to_numpy=True)[0]

                if rag_mode == "チャンク統合":
                    docs = search_similar_documents(query_embedding, tablename)
                    context = "\n\n".join(f"\U0001F4C4 **{d['filename']}**: {d['snippet']}" for d in docs)
                    prompt = f"""
以下の情報を参考にして、質問に答えてください。

情報:
{context}

質問:
{query}
""".strip()
                    response = LLM.invoke(prompt)
                    elapsed_time = round(time.time() - start_time, 2)

                    st.markdown(f"### \U0001F9E0 回答（{elapsed_time} 秒）")
                    st.write(response)

                    st.session_state.history.append({
                        "query": query,
                        "response": response,
                        "model": selected_label,
                        "time": elapsed_time
                    })

                    st.markdown("### \U0001F4D1 検索結果プレビュー")
                    for d in docs:
                        st.markdown(f"**{d['filename']}**")
                        st.write(d["snippet"])
                        with st.expander("▶️ 全文をプレビュー"):
                            unique_key = make_unique_key(d['file_id'], d['filename'])
                            if st.button("✏️ このファイルを編集する", key=f"gotoedit_{unique_key}"):
                                st.session_state.edit_target_file_id = d["file_id"]
                                st.session_state.mode = "ファイル編集"
                                st.rerun()
                            edited_text = st.text_area("全文テキスト（編集可能）", value=d["full_text"], height=200, key=f"edit_{unique_key}")
                            if st.button("保存して再ベクトル化", key=f"save_{unique_key}"):
                                try:
                                    with engine.begin() as conn:
                                        conn.execute(text("UPDATE files SET content = :content WHERE file_id = :file_id"),
                                                     {"content": edited_text, "file_id": d["file_id"]})
                                        conn.execute(text(f'DELETE FROM "{tablename}" WHERE file_id = :file_id'),
                                                     {"file_id": d["file_id"]})
                                    st.success("✅ contentを更新し、旧ベクトルを削除しました")
                                    embed_and_insert(texts=[edited_text], filename=d["filename"], truncate_done_tables=set())
                                    st.success("✅ 再ベクトル化完了！")
                                except Exception as e:
                                    st.error(f"❌ エラーが発生しました: {e}")

                            st.download_button(
                                label="元ファイルをダウンロード",
                                data=bytes(d["file_blob"]),
                                file_name=d["filename"],
                                mime="application/octet-stream",
                                key=f"download_{unique_key}"
                            )

                else:
                    embedding_str = to_pgvector_literal(query_embedding)
                    sql = f"""
                        SELECT DISTINCT f.file_id, f.filename, f.content,
                        MIN(e.embedding <-> '{embedding_str}'::vector) AS distance
                        FROM "{tablename}" AS e
                        JOIN files AS f ON e.file_id = f.file_id
                        GROUP BY f.file_id, f.filename, f.content
                        ORDER BY distance ASC
                        LIMIT 10
                    """
                    with engine.connect() as conn:
                        # REM: SQLAlchemy の Row → dict 変換には mappings() を使う
                        rows = conn.execute(text(sql)).mappings().all()

                    summaries = []
                    for row in rows:
                        summary, score = llm_summarize_with_score(query, row["content"])
                        summaries.append({"file_id": row["file_id"], "filename": row["filename"], "summary": summary, "score": score})

                    summaries.sort(key=lambda x: x["score"], reverse=True)

                    st.markdown("### \U0001F4D1 ファイル別の結果（スコア順）")
                    for s in summaries:
                        st.markdown(f"**📄 {s['filename']}（一致度: {s['score']:.2f}）**")
                        st.write(s["summary"])

    with right_col:
        st.markdown("### 📜 検索履歴")
        if st.session_state.history:
            for i, item in enumerate(reversed(st.session_state.history), 1):
                with st.expander(f"{i}. 質問: {item['query']}（モデル: {item['model']}）", expanded=False):
                    st.markdown(f"🧠 回答（{item['time']} 秒）:\n\n{item['response']}")
        else:
            st.write("まだ履歴はありません。")

        st.markdown("### 💾 履歴をファイルに保存")
        format_type = st.selectbox("保存形式を選択してください", ["txt", "rtf", "docx"])
        if st.button("📥 ダウンロード"):
            file_buffer = generate_history_file(format_type)
            mime_map = {
                "txt": "text/plain",
                "rtf": "application/rtf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            today_str = datetime.date.today().strftime("%Y-%m-%d")
            st.download_button("ファイルをダウンロード", data=file_buffer, file_name=f"search_history_{today_str}.{format_type}", mime=mime_map[format_type])
